"""Parse rows from form's response."""
import logging
import os
from tempfile import NamedTemporaryFile
from time import sleep
from urllib.parse import urlencode

import firebase_admin
import google.auth
import pandas as pd
from docx import Document
from dotenv import load_dotenv
from firebase_admin import credentials, firestore
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaFileUpload
from pandas import Series

logger = logging.getLogger(__name__)

LOGGING_LEVEL = logging.DEBUG
logger.setLevel(LOGGING_LEVEL)

ch = logging.StreamHandler()
ch.setLevel(LOGGING_LEVEL)

formatter = logging.Formatter("%(message)s")
ch.setFormatter(formatter)
logger.addHandler(ch)


class App:
    """Base app."""

    def __init__(self) -> None:
        self.config = self._load_config()
        self._init_storage()

        # Authenticating to Google API.
        gapi_cred, project = google.auth.default()
        self.gservice = build("drive", "v3", credentials=gapi_cred)

    def _init_storage(self) -> None:
        """Configure the working storage."""
        if self.config["working_storage"]["type"] == "firebase":
            firebase_cred = credentials.Certificate(
                self.config["working_storage"]["firebase"]["cred"]
            )
            firebase_admin.initialize_app(
                firebase_cred,
                {"projectId": self.config["working_storage"]["firebase"]["project_id"]},
            )
            self.db = firestore.client()

    def _load_config(self) -> dict:
        """Load the configuration."""
        config = {}
        config["form_response_id"] = os.getenv("APP_FORM_RESPONSE_ID")
        config["sheet_name"] = os.getenv("APP_SHEET_NAME")

        working_storage_type = os.getenv("WORKING_STORAGE_TYPE", "file")
        config["working_storage"] = {}
        config["working_storage"]["type"] = working_storage_type
        if working_storage_type == "firebase":
            logger.info(f"Working storage type: {working_storage_type}")
            config["working_storage"]["firebase"] = {}
            config["working_storage"]["firebase"]["cred"] = os.getenv(
                "GOOGLE_APPLICATION_CREDENTIALS"
            )
            config["working_storage"]["firebase"]["project_id"] = os.getenv(
                "FIREBASE_PROJECT_ID"
            )
            config["working_storage"]["firebase"]["collection_name"] = os.getenv(
                "WORKING_COLLECTION_NAME"
            )
            config["working_storage"]["firebase"]["document_name"] = os.getenv(
                "WORKING_DOCUMENT_NAME"
            )

        return config

    def _create_csv_url(self, sheet_id: str, sheet_name: str) -> str:
        """Create Google's Drive CSV download url."""
        query = {"tqx": "out:csv", "sheet": sheet_name}
        qs = urlencode(query)
        return f"https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?{qs}"

    def get_last_processed_row(self) -> int:
        """Retrieve the last processed row index."""
        if self.config["working_storage"]["type"] == "firebase":
            collection_name = self.config["working_storage"]["firebase"][
                "collection_name"
            ]
            document_name = self.config["working_storage"]["firebase"]["document_name"]
            logger.debug(f"Getting {collection_name}.{document_name} from firebase")

            doc_ref = self.db.collection(collection_name).document(document_name).get()
            doc = doc_ref.to_dict()
            return doc["last_processed_row"]

    def update_last_processed_row(self, row_idx: int) -> None:
        """Write last processed row index back to the storage."""
        if self.config["working_storage"]["type"] == "firebase":
            collection_name = self.config["working_storage"]["firebase"][
                "collection_name"
            ]
            document_name = self.config["working_storage"]["firebase"]["document_name"]
            self.db.collection(collection_name).document(document_name).set(
                {"last_processed_row": row_idx}
            )

    def parse(self) -> None:
        """Parse the rows from the response sheet."""
        # Get the responses.
        df = pd.read_csv(
            self._create_csv_url(
                self.config["form_response_id"], self.config["sheet_name"]
            )
        )
        last_processed_row = self.get_last_processed_row()
        logger.debug(f"Columns of the responses: {df.columns.values.tolist()}")

        # E203 of flake8 is fighting with black.
        to_be_processed = df.iloc[last_processed_row + 1 :]  # noqa: E203
        logger.info(f"Found {to_be_processed.shape[0]} new rows")

        idx = last_processed_row
        for idx, row in to_be_processed.iterrows():
            self.parse_row(idx, row)
            sleep(3)
            last_processed_row = idx
        self.update_last_processed_row(last_processed_row)


class COTS2021(App):
    """Implement parsing for COTS 2021."""

    def __init__(self):
        super().__init__()
        logger.debug("Initializing COTS2021")
        self.config["folder_ids"] = {}
        self.config["folder_ids"]["individual"] = os.getenv(
            "INDIVIDUAL_PROPOSAL_FOLDER_ID"
        )
        self.config["folder_ids"]["panel"] = os.getenv("PANEL_PROPOSAL_FOLDER_ID")
        self.config["folder_ids"]["roundtable"] = os.getenv(
            "ROUNDTABLE_PROPOSAL_FOLDER_ID"
        )

    def create_base_doc(self, idx: int, row: Series, heading: str) -> Document:
        """Create base document with common information."""
        fullname = row["First name"] + " " + row["Last name"]

        document = Document()
        normal_styles = document.styles["Normal"]
        normal_styles.font.name = "Arial"
        document.add_heading(heading, 0)

        p = document.add_paragraph("")
        p.add_run("Submitted by:").bold = True
        p.add_run(f" {fullname} ({row['Email address']})")
        p.add_run(" At:").bold = True
        p.add_run(f" {row['Timestamp']}")

        p = document.add_paragraph("")
        p.add_run("Affiliation:").bold = True
        p.add_run(f" {row['Affiliation']}")
        return document

    def add_metadata(self, idx: int, row: Series, document: Document) -> None:
        """Add metadata to the document."""
        pass

    def create_individual_proposal_doc(self, idx: int, row: Series) -> Document:
        """Create document for individual proposal."""
        document = self.create_base_doc(idx, row, "Individual/Film/Other")

        document.add_heading("Abstract", 1)
        document.add_paragraph(row["Abstract"])
        return document

    def create_panel_proposal_doc(self, idx: int, row: Series) -> Document:
        """Create document for panel proposal."""
        document = self.create_base_doc(idx, row, "Panel")

        document.add_heading("Topic", 1)
        document.add_paragraph(row["Topic of the panel"])

        document.add_heading("Panelists", 1)
        document.add_paragraph(row["Names of the panelists"])

        document.add_heading("Emails", 1)
        document.add_paragraph(row["Contact information of the panelists"])

        document.add_heading("Abstracts", 1)
        document.add_paragraph(row["Abstracts"])
        return document

    def create_roundtable_proposal_doc(self, idx: int, row: Series) -> Document:
        """Create document for roundtable proposal."""
        document = self.create_base_doc(idx, row, "Roundtable")

        document.add_heading("Abstract", 1)
        document.add_paragraph(row["Abstract.1"])

        document.add_heading("Participants", 1)
        document.add_paragraph(row["Name of the participants"])

        document.add_heading("Emails", 1)
        document.add_paragraph(row["Contact information of the participants"])
        return document

    def parse_row(self, idx: int, row: Series) -> None:
        """Parse row into a document."""
        fullname = row["First name"] + " " + row["Last name"]
        with NamedTemporaryFile() as temp_file:
            if (
                row["What proposal are you planning to submit"]
                == "A proposal for an individual paper, film screening or other presentationtation"
            ):
                parent_folder_id = self.config["folder_ids"]["individual"]
                proposal_type = "individual"
                document = self.create_individual_proposal_doc(idx, row)
            elif (
                row["What proposal are you planning to submit"]
                == "A proposal for a paper panel"
            ):
                parent_folder_id = self.config["folder_ids"]["panel"]
                proposal_type = "panel"
                document = self.create_panel_proposal_doc(idx, row)
            elif (
                row["What proposal are you planning to submit"]
                == "A proposal for a routable"
            ):
                parent_folder_id = self.config["folder_ids"]["roundtable"]
                proposal_type = "roundtable"
                document = self.create_roundtable_proposal_doc(idx, row)

            # Create the word document.
            document.save(temp_file.name)

            # Upload the document.
            gdoc_name = f"{fullname} (id-{idx} type-{proposal_type})"
            file_metadata = {
                "name": gdoc_name,
                "mimeType": "application/vnd.google-apps.document",
                "parents": [parent_folder_id],
            }
            media = MediaFileUpload(
                temp_file.name,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

            try:
                logger.info(
                    "Uploading {} as {} to {}".format(
                        str(temp_file.name), gdoc_name, parent_folder_id
                    )
                )
                self.gservice.files().create(
                    body=file_metadata, media_body=media, fields="id"
                ).execute()
            except HttpError:
                logger.error("HttpError occurred during an upload")


if __name__ == "__main__":
    load_dotenv()
    app = COTS2021()
    app.parse()
